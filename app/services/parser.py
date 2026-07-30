"""Unified Resume Parsing Service (PDF, DOCX, TXT/MD).

Supports parsing raw bytes or file paths into a structured ParsedResume object.
Extracts section-by-section details (Name, Email, Phone, Summary, Skills, Experience, Education, Projects).
Validates file extensions, size limits, and sanitizes input data without inserting fake emails.
"""
from __future__ import annotations

import io
import logging
import os
import re
from typing import Any
from pydantic import BaseModel, Field

logger = logging.getLogger("talentops.parser")

_EMAIL_REGEX = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
_PHONE_REGEX = re.compile(r"\+?\d{1,3}[-.\s]?(?:\(\d{1,4}\)[-.\s]?)?\d{2,5}[-.\s]?\d{3,5}")
_ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
_DEFAULT_MAX_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB limit

TECH_SKILLS_KEYWORDS = [
    "python", "java", "c++", "golang", "rust", "javascript", "typescript",
    "react", "vue", "angular", "node.js", "express", "django", "fastapi", "flask",
    "postgresql", "postgres", "mysql", "mongodb", "redis", "docker", "kubernetes",
    "aws", "gcp", "azure", "asyncio", "graphql", "rest api", "restful", "ci/cd", "git",
    "system design", "microservices", "kafka", "pytorch", "tensorflow"
]


class ResumeParseError(Exception):
    """Raised when parsing resume file content fails."""
    pass


class UnsupportedFileTypeError(ResumeParseError):
    """Raised when the uploaded file type is not supported."""
    pass


class FileTooLargeError(ResumeParseError):
    """Raised when file size exceeds maximum permitted limit."""
    pass


class CandidateProject(BaseModel):
    """Structured project entry from candidate resume."""
    title: str
    description: str = ""
    technologies: list[str] = Field(default_factory=list)
    url: str = ""


class CandidateExperience(BaseModel):
    """Structured work experience entry."""
    company: str = ""
    role: str = ""
    dates: str = ""
    description: str = ""


class CandidateEducation(BaseModel):
    """Structured education entry."""
    degree: str = ""
    institution: str = ""
    year: str = ""


class ParsedResume(BaseModel):
    """Structured result of resume parsing."""
    raw_text: str
    file_name: str
    file_type: str
    email: str = ""
    candidate_name: str = ""
    phone: str = ""
    summary: str = ""
    skills: list[str] = Field(default_factory=list)
    projects: list[CandidateProject] = Field(default_factory=list)
    experience: list[CandidateExperience] = Field(default_factory=list)
    education: list[CandidateEducation] = Field(default_factory=list)
    metadata: dict[str, Any] = Field(default_factory=dict)


def clean_candidate_name(raw_name: str) -> str:
    """Clean up a raw filename or text line to extract a candidate full name."""
    if not raw_name:
        return ""

    # 1. Remove file extensions
    cleaned = re.sub(r"\.(pdf|docx|doc|txt|md)$", "", str(raw_name).strip(), flags=re.IGNORECASE)

    # 2. Remove UUID prefixes (hex 32 or standard 36-char uuid followed by _ or -)
    cleaned = re.sub(r"^[a-fA-F0-9]{32}[_-]?", "", cleaned)
    cleaned = re.sub(r"^[a-fA-F0-9]{8}-[a-fA-F0-9]{4}-[a-fA-F0-9]{4}-[a-fA-F0-9]{4}-[a-fA-F0-9]{12}[_-]?", "", cleaned)

    # 3. Replace underscores/dashes with spaces first
    cleaned = cleaned.replace("_", " ").replace("-", " ")

    # 4. Remove common junk keywords
    junk_patterns = [
        r"\bAI[ -]?RESUME\b", r"\bRESUME\b", r"\bCV\b", r"\bCURRICULUM\b", r"\bVITAE\b",
        r"\bSingle[ -]?P(age)?\b", r"\bDraft\b", r"\bFinal\b", r"\bCopy\b", r"\bUpload\b",
        r"\bDocument\b", r"\bProfile\b"
    ]
    for pat in junk_patterns:
        cleaned = re.sub(pat, "", cleaned, flags=re.IGNORECASE)

    # 5. Strip non-alpha characters and collapse whitespace
    cleaned = re.sub(r"[^a-zA-Z\s\.]", "", cleaned)
    cleaned = " ".join(cleaned.split())

    if cleaned and len(cleaned) >= 2:
        return cleaned.title()
    return ""


def extract_email_from_text(text: str) -> str:
    """Extract candidate email address from resume text. Never invents fake emails."""
    matches = _EMAIL_REGEX.findall(text or "")
    if matches:
        for m in matches:
            if not any(ignore in m.lower() for ignore in ["example.com", "domain.com", "github.com", "w3.org"]):
                return m
        return matches[0]
    return ""


def extract_phone_from_text(text: str) -> str:
    """Extract phone number from resume text using regex."""
    top_lines = (text or "").splitlines()[:20]
    top_text = "\n".join(top_lines)
    matches = _PHONE_REGEX.findall(top_text)
    if matches:
        for m in matches:
            digits = re.sub(r"\D", "", m)
            if 7 <= len(digits) <= 15:
                return m.strip()
    return ""


def extract_skills_word_boundary(text: str) -> list[str]:
    """Extract skills using word boundaries to prevent false positives (e.g. 'go' in 'going')."""
    lower_text = (text or "").lower()
    found_skills = []
    for skill in TECH_SKILLS_KEYWORDS:
        # Escape special regex chars like ++ in C++
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, lower_text):
            found_skills.append(skill)
    return found_skills


def extract_candidate_metadata(resume_text: str, file_name: str | None = None) -> dict[str, str]:
    """Extract candidate full name, email, and phone from resume text.
    
    Strictly avoids fake email generation (@example.com) and fallback defaults like 'Candidate'.
    """
    email = extract_email_from_text(resume_text or "")
    phone = extract_phone_from_text(resume_text or "")

    extracted_name = ""
    lines = [line.strip() for line in (resume_text or "").splitlines() if line.strip()]
    top_lines = lines[:15]

    ignore_words = {
        "curriculum", "vitae", "resume", "cv", "summary", "experience", "education",
        "profile", "contact", "page", "phone", "email", "skills", "projects", "senior",
        "junior", "lead", "staff", "principal", "engineer", "developer", "architect",
        "manager", "data", "software", "fullstack", "backend", "frontend", "technologies",
        "objective", "work", "history", "name", "candidate", "applicant"
    }

    for line in top_lines:
        if "@" in line or "http" in line or "www." in line or "linkedin" in line or "github" in line:
            continue
        if re.search(r"\+?\d[\d\s-]{7,}", line):
            continue

        # Strip markdown symbols and explicit name prefixes
        clean_line = re.sub(r"[#\*\_\`]", "", line).strip()
        clean_line = re.sub(r"^(?:full\s+name|candidate\s+name|applicant\s+name|name|applicant|candidate)\s*[:\-]\s*", "", clean_line, flags=re.IGNORECASE).strip()

        segments = [s.strip() for s in re.split(r"\s*[-|\:\,]\s*", clean_line) if s.strip()]
        for segment in segments:
            words = segment.split()
            if 1 <= len(words) <= 4:
                clean_words = [re.sub(r"[^a-zA-Z]", "", w) for w in words]
                clean_words = [w for w in clean_words if w]
                if clean_words and not any(w.lower() in ignore_words for w in clean_words):
                    candidate_cand = " ".join(clean_words)
                    if len(candidate_cand) >= 2:
                        extracted_name = candidate_cand.title()
                        break
        if extracted_name:
            break

    # If line parsing didn't find a name, try inferring from email handle (e.g. john.doe@email.com -> John Doe)
    if not extracted_name and email:
        handle = email.split("@")[0]
        # Only use handle if it contains separators or distinct words (not random hex/hash)
        parts = re.split(r"[._\-+]", handle)
        clean_parts = [re.sub(r"[^a-zA-Z]", "", p) for p in parts]
        clean_parts = [p for p in clean_parts if len(p) >= 2 and p.lower() not in ignore_words]
        if 1 <= len(clean_parts) <= 3:
            extracted_name = " ".join(clean_parts).title()

    if not extracted_name and file_name:
        extracted_name = clean_candidate_name(file_name)
    elif extracted_name:
        extracted_name = clean_candidate_name(extracted_name)

    return {
        "full_name": extracted_name,
        "email": email,
        "phone": phone,
    }


def extract_sections_by_regex(text: str) -> dict[str, str]:
    """Parse text into section headers: SUMMARY, SKILLS, PROJECTS, EXPERIENCE, EDUCATION."""
    headers_pat = re.compile(
        r"\n(?=\s*(?:(?:PERSONAL\s+|ACADEMIC\s+)?PROJECTS?|(?:WORK\s+|PROFESSIONAL\s+)?EXPERIENCE|EMPLOYMENT\s+HISTORY|EDUCATION|ACADEMICS?|TECHNICAL\s+SKILLS|SKILLS|SUMMARY|PROFESSIONAL\s+SUMMARY|OBJECTIVE)\b[:\s\n])",
        re.IGNORECASE,
    )
    parts = headers_pat.split(text)
    sections: dict[str, str] = {}

    for part in parts:
        part_clean = part.strip()
        if not part_clean:
            continue
        header_match = re.match(
            r"^(?:PERSONAL\s+|ACADEMIC\s+)?(PROJECTS?|(?:WORK\s+|PROFESSIONAL\s+)?EXPERIENCE|EMPLOYMENT\s+HISTORY|EDUCATION|ACADEMICS?|TECHNICAL\s+SKILLS|SKILLS|SUMMARY|PROFESSIONAL\s+SUMMARY|OBJECTIVE)\b[:\s\n]*(.*)",
            part_clean,
            re.IGNORECASE | re.DOTALL,
        )
        if header_match:
            hdr_name = header_match.group(1).upper()
            hdr_body = header_match.group(2).strip()
            if "PROJECT" in hdr_name:
                sections["projects"] = hdr_body
            elif "EXPERIENCE" in hdr_name or "EMPLOYMENT" in hdr_name:
                sections["experience"] = hdr_body
            elif "EDUCATION" in hdr_name or "ACADEMIC" in hdr_name:
                sections["education"] = hdr_body
            elif "SKILL" in hdr_name:
                sections["skills"] = hdr_body
            elif "SUMMARY" in hdr_name or "OBJECTIVE" in hdr_name:
                sections["summary"] = hdr_body
        else:
            if "header" not in sections:
                sections["header"] = part_clean

    return sections


def extract_projects_from_section(projects_text: str) -> list[CandidateProject]:
    """Extract individual projects from a project section block."""
    if not projects_text or not projects_text.strip():
        return []

    projects: list[CandidateProject] = []
    # Split by bullet headers or blank lines/numeric prefixes
    items = re.split(r"\n(?=\s*(?:[•\-*\d+\.]|[A-Z][a-zA-Z0-9\s]{3,30}:))", projects_text)

    for item in items:
        item_str = item.strip()
        if not item_str or len(item_str) < 5:
            continue
        lines = [l.strip() for l in item_str.splitlines() if l.strip()]
        title = lines[0].lstrip("•-*123456789. ").strip()
        desc = " ".join(lines[1:]) if len(lines) > 1 else ""

        # Extract URLs if any
        urls = re.findall(r"https?://[^\s]+", item_str)
        url = urls[0] if urls else ""

        # Extract tech stack
        skills_used = extract_skills_word_boundary(item_str)

        if title:
            projects.append(
                CandidateProject(
                    title=title[:150],
                    description=desc[:1000],
                    technologies=skills_used,
                    url=url,
                )
            )

    return projects


def parse_pdf_bytes(pdf_bytes: bytes) -> str:
    """Extract text from raw PDF bytes using pypdf with fallback to pdfplumber."""
    for pypdf_log_name in ["pypdf", "pypdf._reader", "pypdf.filters", "pypdf.generic._data_structures"]:
        logging.getLogger(pypdf_log_name).setLevel(logging.ERROR)

    is_pdf = pdf_bytes.startswith(b"%PDF") or b"%PDF-" in pdf_bytes[:1024]

    if is_pdf:
        extracted = ""
        # 1. Primary parser: pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(pdf_bytes))
            pages = [(page.extract_text() or "") for page in reader.pages]
            extracted = "\n".join(t for t in pages if t.strip()).strip()
        except Exception as e:
            logger.warning("pypdf failed to extract text from PDF: %s", e)

        # 2. Secondary fallback: pdfplumber
        if not extracted or len(extracted) < 10:
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                    pages = [(page.extract_text() or "") for page in pdf.pages]
                    extracted = "\n".join(t for t in pages if t.strip()).strip()
            except Exception as e:
                logger.warning("pdfplumber failed to extract text from PDF: %s", e)

        if not extracted or not extracted.strip():
            raise ResumeParseError("Could not extract text from PDF. Ensure PDF is not scanned/image-only.")

        return extracted

    # Fallback ONLY if file is a non-PDF plain text file
    try:
        text = pdf_bytes.decode("utf-8", errors="strict")
        if text.strip():
            logger.info("Parsed plain text file without PDF header")
            return text
    except Exception:
        pass

    raise ResumeParseError("Failed to parse PDF content: Invalid or corrupt PDF binary structure")


def parse_docx_bytes(docx_bytes: bytes) -> str:
    """Extract text from raw DOCX bytes using python-docx or zipfile XML fallback."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(docx_bytes))
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += "\n" + row_text
        return text
    except ImportError:
        logger.warning("python-docx not installed; attempting XML extraction fallback")
        try:
            import zipfile
            from xml.etree import ElementTree as ET
            with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
                xml_content = zf.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                paragraphs = []
                for p in tree.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                    texts = [t.text for t in p.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t") if t.text]
                    if texts:
                        paragraphs.append("".join(texts))
                return "\n".join(paragraphs)
        except Exception as ex:
            raise ResumeParseError(f"Failed to parse DOCX bytes via XML fallback: {ex}") from ex
    except Exception as e:
        logger.error("Failed to parse DOCX bytes: %s", e)
        raise ResumeParseError(f"Failed to parse DOCX content: {e}") from e


async def parse_resume_bytes(
    content: bytes,
    file_name: str = "resume.pdf",
    max_size_bytes: int = _DEFAULT_MAX_SIZE_BYTES,
) -> ParsedResume:
    """Parse resume raw bytes into a structured ParsedResume object with section-by-section breakdown."""
    if len(content) > max_size_bytes:
        raise FileTooLargeError(
            f"File size ({len(content)} bytes) exceeds maximum limit of {max_size_bytes} bytes"
        )

    ext = os.path.splitext(file_name)[1].lower() or ".pdf"
    if ext not in _ALLOWED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"File extension '{ext}' is not supported. Supported extensions: {', '.join(sorted(_ALLOWED_EXTENSIONS))}"
        )

    file_type = ext.lstrip(".")

    if file_type == "pdf":
        raw_text = parse_pdf_bytes(content)
    elif file_type == "docx":
        raw_text = parse_docx_bytes(content)
    else:  # txt or md
        try:
            raw_text = content.decode("utf-8", errors="replace")
        except Exception as e:
            raise ResumeParseError(f"Failed to decode text file: {e}") from e

    if not raw_text or not raw_text.strip():
        raise ResumeParseError("Extracted text from resume is empty or missing")

    meta = extract_candidate_metadata(raw_text, file_name=file_name)
    email = meta.get("email") or ""
    phone = meta.get("phone") or ""
    candidate_name = meta.get("full_name") or ""

    # Call the new LLM-based parser agent
    from app.agents.parser_agent import parse_resume_with_llm
    
    llm_data = {}
    try:
        llm_data = await parse_resume_with_llm(raw_text)
    except Exception as e:
        logger.error("LLM parser failed: %s, falling back to regex extraction", e)

    # Use LLM data if present, otherwise fallback to regex extraction
    if llm_data:
        summary = llm_data.get("summary") or ""
        skills = llm_data.get("skills") or []
        
        projects = []
        for p in llm_data.get("projects", []):
            projects.append(CandidateProject(
                title=p.get("title", "Unknown Project")[:150],
                description=p.get("description", "")[:1000],
                technologies=p.get("technologies", []),
                url=p.get("url", "")
            ))
            
        experience = []
        for e in llm_data.get("experience", []):
            experience.append(CandidateExperience(
                company=e.get("company", ""),
                role=e.get("role", ""),
                dates=e.get("dates", ""),
                description=e.get("description", "")[:1000]
            ))
            
        education = []
        for edu in llm_data.get("education", []):
            education.append(CandidateEducation(
                degree=edu.get("degree", "")[:500],
                institution=edu.get("institution", ""),
                year=edu.get("year", "")
            ))
    else:
        # Regex Fallback
        sections = extract_sections_by_regex(raw_text)
        summary = sections.get("summary") or ""
        skills = extract_skills_word_boundary(raw_text)

        # Extract projects from projects section or full text
        projects_raw = sections.get("projects") or ""
        projects = extract_projects_from_section(projects_raw)
        
        if not projects:
            # Fallback: look for generic project indicators in the raw text
            proj_matches = re.findall(r"(?:^|\n)([^\n]*?(?:github\.com|built a|developed a|created a|personal project)[^\n]*(?:\n[^\n]*){0,3})", raw_text, re.IGNORECASE)
            if proj_matches:
                for match in set(proj_matches):
                    if len(match.strip()) > 20:
                        projects.append(CandidateProject(
                            title="Inferred Project",
                            description=match.strip()[:1000],
                            technologies=extract_skills_word_boundary(match)
                        ))
                        break # just grab one for fallback

        experience_raw = sections.get("experience") or ""
        experience = []
        if experience_raw.strip():
            experience.append(CandidateExperience(description=experience_raw.strip()[:1000]))

        education_raw = sections.get("education") or ""
        education = []
        if education_raw.strip():
            education.append(CandidateEducation(degree=education_raw.strip()[:500]))
        
        if not education:
            # Fallback for education
            edu_matches = re.findall(r"(?:^|\n)([^\n]*?(?:University|College|Bachelor|Master|PhD|B\.S\.|B\.Tech|M\.Tech|M\.S\.|B\.A\.|M\.A\.)[^\n]*)", raw_text, re.IGNORECASE)
            if edu_matches:
                edu_text = " | ".join(set(edu_matches))[:500]
                if edu_text.strip():
                    education.append(CandidateEducation(degree=edu_text.strip()))

    return ParsedResume(
        raw_text=raw_text,
        file_name=file_name,
        file_type=file_type,
        email=email,
        candidate_name=candidate_name,
        phone=phone,
        summary=summary,
        skills=skills,
        projects=projects,
        experience=experience,
        education=education,
        metadata={"content_length": len(content), "char_count": len(raw_text)}
    )


async def parse_resume(path: str) -> ParsedResume:
    """Parse a resume file path into ParsedResume."""
    if not os.path.exists(path):
        raise ResumeParseError(f"Resume file path does not exist: {path}")

    with open(path, "rb") as f:
        content = f.read()

    filename = os.path.basename(path)
    return await parse_resume_bytes(content, file_name=filename)
